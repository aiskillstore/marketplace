#!/usr/bin/env python3
"""Validate SOW/PWS DOCX structure, headings, TOC, and artifact separation."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


CORE_SECTIONS = [
    "introduction",
    "definitions and acronyms",
    "requirements",
    "deliverables",
    "period of performance",
    "place of performance",
    "government-furnished property and information",
    "security and privacy",
    "key personnel",
    "reporting and oversight",
    "quality",
    "transition",
    "constraints and assumptions",
]

FORBIDDEN = {
    "staffing handoff": re.compile(r"staffing\s+handoff\s+table", re.I),
    "IGCE skill plumbing": re.compile(r"\bIGCE\s+Builder\b|\bbuild\s+the\s+IGCE\b", re.I),
    "Section B handoff": re.compile(r"(?:CLIN|Section\s+B)\s+handoff\s+table", re.I),
    "SOC code": re.compile(r"\bSOC\s+Code\b", re.I),
    "FTE staffing": re.compile(r"\bFTEs?\b", re.I),
    "CLIN content": re.compile(r"\bCLINs?\b", re.I),
    "labor-category ceiling-hours table": re.compile(r"Labor\s+Category\s+Ceiling\s+Hours", re.I),
    "wrong FAR 37 citation": re.compile(r"FAR\s+37\.102\s*\(d\)", re.I),
    "wrong key-personnel clause": re.compile(r"FAR\s+52\.237-2", re.I),
    "TOC refresh instruction": re.compile(r"(?:Ctrl|Cmd)\+A.{0,40}F9", re.I | re.S),
    "local runtime path": re.compile(r"/(?:mnt|tmp|Users)/|[A-Za-z]:\\", re.I),
}


def normalize_heading(text: str) -> str:
    text = re.sub(r"^\s*(?:section\s+)?\d+(?:\.\d+)*[.):-]?\s*", "", text, flags=re.I)
    return re.sub(r"\s+", " ", text).strip().lower()


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def heading_level(paragraph: object) -> int | None:
    style = getattr(paragraph, "style", None)
    name = getattr(style, "name", "") or ""
    match = re.fullmatch(r"Heading\s+([1-9])", name, re.I)
    return int(match.group(1)) if match else None


def has_toc_and_update(path: Path) -> tuple[bool, bool]:
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        settings_xml = archive.read("word/settings.xml").decode("utf-8", errors="replace")
    has_toc = bool(re.search(r"<w:instrText[^>]*>[^<]*\bTOC\b", document_xml, re.I))
    update = bool(re.search(r"<w:updateFields\b[^>]*(?:w:val=[\"'](?:true|1)[\"'])?", settings_xml, re.I))
    return has_toc, update


def validate(path: Path, document_type: str) -> dict[str, object]:
    failures: list[str] = []
    try:
        if not zipfile.is_zipfile(path):
            return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            for required in ("[Content_Types].xml", "word/document.xml", "word/styles.xml"):
                if required not in names:
                    failures.append(f"missing DOCX part: {required}")
        document = Document(path)
    except (OSError, ValueError, zipfile.BadZipFile) as exc:
        return {"status": "fail", "failures": [f"cannot read DOCX: {exc}"]}

    text = document_text(document)
    for label, pattern in FORBIDDEN.items():
        if pattern.search(text):
            failures.append(f"document contains forbidden {label}")

    h1 = [
        normalize_heading(paragraph.text)
        for paragraph in document.paragraphs
        if heading_level(paragraph) == 1 and paragraph.text.strip()
    ]
    if len(h1) < len(CORE_SECTIONS):
        failures.append(f"only {len(h1)} Heading 1 sections found; expected at least {len(CORE_SECTIONS)}")

    positions: list[int] = []
    for expected in CORE_SECTIONS:
        if expected == "quality":
            candidates = [
                index
                for index, value in enumerate(h1)
                if value in {"qasp summary", "inspection and acceptance"}
            ]
        else:
            candidates = [index for index, value in enumerate(h1) if value == expected]
        if not candidates:
            failures.append(f"missing Heading 1 section: {expected}")
        else:
            positions.append(candidates[0])
    if positions and positions != sorted(positions):
        failures.append("core Heading 1 sections are out of order")

    if document_type == "pws":
        if "qasp summary" not in h1:
            failures.append("PWS is missing a Heading 1 QASP Summary")
        if not re.search(r"\bAQL\b|Acceptable Quality Level", text, re.I):
            failures.append("PWS does not contain an AQL")
        if not re.search(r"performance standard|assessment method|method of assessment", text, re.I):
            failures.append("PWS does not contain measurable performance or assessment language")
    else:
        if "inspection and acceptance" not in h1:
            failures.append("SOW is missing a Heading 1 Inspection and Acceptance section")

    if re.search(r"\bCPARS\b", text, re.I) and re.search(
        r"\b(?:Exceptional|Very Good|Satisfactory|Marginal|Unsatisfactory)\b", text, re.I
    ):
        failures.append("document ties CPARS language to rating labels")

    if len(h1) > 8:
        try:
            has_toc, update = has_toc_and_update(path)
        except (KeyError, OSError, zipfile.BadZipFile) as exc:
            failures.append(f"cannot audit TOC settings: {exc}")
        else:
            if not has_toc:
                failures.append("document with more than eight sections has no dynamic TOC field")
            if not update:
                failures.append("document does not set updateFields-on-open")

    return {
        "status": "pass" if not failures else "fail",
        "document_type": document_type,
        "heading_1_count": len(h1),
        "table_count": len(document.tables),
        "failures": failures,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate a SOW or PWS DOCX.")
    parser.add_argument("document", type=Path)
    parser.add_argument("--document-type", required=True, choices=("sow", "pws"))
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    if not args.document.is_file():
        print(f"ERROR: document not found: {args.document}", file=sys.stderr)
        return 2
    result = validate(args.document, args.document_type)
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("DOCX structure, content separation, and TOC settings passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
