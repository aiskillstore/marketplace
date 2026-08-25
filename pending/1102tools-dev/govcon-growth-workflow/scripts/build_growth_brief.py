#!/usr/bin/env python3
"""Build a structured GovCon Growth Brief from a validated research record."""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
GREEN = "167D5A"
GRAY = "5B6573"


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_cell_text(cell, value: object, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(value if value not in (None, "") else "Not provided"))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12.5, GREEN)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    if "Evidence ID" not in document.styles:
        style = document.styles.add_style("Evidence ID", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Aptos"
        style.font.size = Pt(8.5)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN)
    header = section.header.paragraphs[0]
    header.text = "1102tools  |  GovCon Growth"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Prepared as of the date shown  |  Page ").font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float],
    *,
    repeat_header: bool = True,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[index], NAVY)
        table.rows[0].cells[index].width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if row_index % 2:
                shade(cells[index], "F4F6F8")
            cells[index].width = Inches(widths[index])
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    if repeat_header:
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def add_bullets(document: Document, items: list[object], empty: str) -> None:
    if not items:
        document.add_paragraph(empty)
        return
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("decision") or item.get("question") or json.dumps(item, sort_keys=True)
        else:
            text = str(item)
        paragraph = document.add_paragraph(text, style="List Bullet")
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.05


def cite_ids(paragraph, ids: list[str]) -> None:
    if ids:
        paragraph.add_run(" [")
        paragraph.add_run(", ".join(ids), style="Evidence ID")
        paragraph.add_run("]")


def build(record: dict, output: Path) -> None:
    document = Document()
    configure(document)
    as_of = record.get("scope", {}).get("as_of_date", "Not stated")
    has_bid_decision = bool(record.get("validation", {}).get("bid_context_complete"))
    label = "GovCon Growth Brief" if has_bid_decision else "Evidence Brief - No Bid Decision"

    brand = document.add_paragraph()
    brand.paragraph_format.space_before = Pt(115)
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("1102tools")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(GREEN)
    title = document.add_paragraph(label, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(record.get("question", "GovCon growth research"))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    meta = document.add_paragraph(f"As of {as_of}\nPrepared by: ____________________")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    headings = [
        "Executive Summary",
        "Business Question and Scope",
        "Company Context and Missing Inputs",
        "Evidence and Analysis",
        "Assessment or Pipeline",
        "Risks, Contrary Evidence, and Limitations",
        "User Decision and Next Actions",
        "Reproducible Search Log",
        "Evidence Register",
    ]
    document.add_heading(headings[0], level=1)
    document.add_paragraph(record.get("validation", {}).get("executive_summary", "This brief organizes the approved public and internal evidence."))
    if not has_bid_decision:
        quote = document.add_paragraph(style="Intense Quote")
        quote.add_run("Decision boundary: ").bold = True
        quote.add_run("Internal company context is incomplete. This brief contains no bid or no-bid recommendation.")

    document.add_heading(headings[1], level=1)
    document.add_paragraph(record.get("question", "Not provided"))
    add_table(document, ["Scope field", "Value"], [[key.replace("_", " ").title(), value] for key, value in record.get("scope", {}).items()], [2.1, 4.8])

    document.add_heading(headings[2], level=1)
    add_bullets(document, record.get("user_context", []), "No internal company context was supplied.")
    document.add_heading("Assumptions", level=2)
    add_bullets(document, record.get("assumptions", []), "No working assumption was recorded.")
    missing = record.get("validation", {}).get("missing_bid_context", [])
    document.add_heading("Missing bid-screen context", level=2)
    add_bullets(document, missing, "None recorded.")

    document.add_heading(headings[3], level=1)
    for finding in record.get("findings", []):
        p = document.add_paragraph(finding.get("text", ""), style="List Bullet")
        cite_ids(p, finding.get("evidence_ids", []))
    if not record.get("findings"):
        document.add_paragraph("No approved finding was recorded.")
    for index, check in enumerate(record.get("validation", {}).get("numeric_checks", [])):
        total = sum(float(value) for value in check.get("components", []))
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in record.get("evidence", [])
            if isinstance(item, dict)
            and item.get("source_class") == "calculation"
            and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            raise ValueError(
                f"numeric check {index} requires exactly one calculation evidence item whose locator is {locator}"
            )
        paragraph = document.add_paragraph(f"{check.get('label', 'Calculated total')}: {total:,.2f}")
        cite_ids(paragraph, calculation_ids)

    document.add_heading(headings[4], level=1)
    document.add_paragraph(record.get("validation", {}).get("assessment", "No final assessment was approved."))
    pipeline = record.get("validation", {}).get("pipeline", [])
    if pipeline:
        add_table(document, ["Candidate", "Signal", "Timing", "Confidence", "Next validation"], [[p.get("candidate", ""), p.get("signal", ""), p.get("timing", ""), p.get("confidence", ""), p.get("next_validation", "")] for p in pipeline], [1.5, 1.8, 1.1, 0.8, 1.7])

    document.add_heading(headings[5], level=1)
    add_bullets(document, record.get("conflicts", []), "No source conflict was recorded.")
    add_bullets(document, record.get("unresolved_questions", []), "No unresolved question was recorded.")
    for inference in record.get("inferences", []):
        p = document.add_paragraph("Inference: " + inference.get("text", inference.get("reasoning", "")), style="List Bullet")
        cite_ids(p, inference.get("evidence_ids", []))

    document.add_heading(headings[6], level=1)
    add_bullets(document, record.get("user_decisions", []), "No user decision was recorded.")
    add_bullets(document, record.get("validation", {}).get("next_actions", []), "No next action was recorded.")

    document.add_heading(headings[7], level=1)
    queries = record.get("queries", [])
    if queries:
        add_table(document, ["Source / operation", "Sanitized parameters", "Retrieved", "Coverage and limits"], [[q.get("operation", q.get("source", "")), json.dumps(q.get("parameters", {}), sort_keys=True), q.get("retrieved_at", ""), f"{q.get('count', 'n/a')}; {q.get('limitations', '')}"] for q in queries], [1.5, 2.35, 1.25, 2.0])
    else:
        document.add_paragraph("No external query was made.")

    document.add_heading(headings[8], level=1)
    # LibreOffice can position a repeated header above the printable area on a
    # later page of a long fixed-width table. Keep the header on the first page
    # only so every continued evidence row remains fully visible after PDF
    # conversion.
    add_table(
        document,
        ["ID", "Class", "Source", "Fact", "Limitations"],
        [[
            e.get("id", ""),
            e.get("source_class", ""),
            f"{e.get('title', '')}\n{e.get('locator', '')}",
            e.get("fact", ""),
            e.get("limitations", ""),
        ] for e in record.get("evidence", [])],
        [0.55, 0.85, 1.55, 2.5, 1.65],
        repeat_header=False,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("record", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    validator = Path(__file__).with_name("validate_research_record.py")
    result = subprocess.run([sys.executable, str(validator), str(args.record)], check=False)
    if result.returncode:
        return result.returncode
    build(json.loads(args.record.read_text(encoding="utf-8")), args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
