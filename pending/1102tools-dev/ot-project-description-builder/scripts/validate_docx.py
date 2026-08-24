#!/usr/bin/env python3
"""Validate an OT project-description DOCX for structure and separation."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


CORE_SECTIONS = [
    "agreement overview",
    "technical background and current state",
    "project objectives",
    "technical approach by phase",
    "milestone schedule",
    "deliverables",
    "data rights",
    "period of performance",
    "government responsibilities",
    "key personnel",
    "reporting and oversight",
    "constraints and assumptions",
]

FORBIDDEN = {
    "milestone handoff": re.compile(r"MILESTONE\s+HANDOFF\s+TABLE", re.I),
    "skill-routing language": re.compile(
        r"\bOT\s+Cost\s+Analysis\s+skill\b|\bbuild\s+the\s+OT\s+cost\s+analysis\b|\$ot-cost-analysis",
        re.I,
    ),
    "internal workpaper notice": re.compile(r"Internal\s+Government\s+workpaper", re.I),
    "currency amount": re.compile(
        r"(?:\$\s*\d|\b\d[\d,]*(?:\.\d+)?\s*(?:dollars?|USD|million|billion)\b)",
        re.I,
    ),
    "pricing or funding content": re.compile(
        r"\bshould[- ]cost\b|\bfunding\s+profile\b|\bGovernment\s+budget\b|"
        r"\bprice[- ]reasonableness\b|\blabor\s+rate\b|\bhourly\s+rate\b|\bmilestone\s+amount\b",
        re.I,
    ),
    "FAR artifact language": re.compile(r"\bCLINs?\b|\bQASP\b|\bAQL\b|\boption\s+year\b", re.I),
    "wrong prototype authority": re.compile(
        r"(?:Prototype\s+OT|prototype\s+(?:project|authority))[^.\n]{0,120}"
        r"(?:under|pursuant\s+to|authorized\s+by|authority:?\s*)\s*10\s*U\.?S\.?C\.?\s*4021",
        re.I,
    ),
    "obsolete prototype-definition citation": re.compile(r"10\s*U\.?S\.?C\.?\s*4003", re.I),
    "false Path D condition": re.compile(
        r"(?:4022\s*\(d\)\s*\(1\)\s*\(D\)|Path\s+D)[^.\n]{0,140}"
        r"(?:competition\s+commitment|commit(?:ment|ted)\s+to\s+compet)",
        re.I,
    ),
    "local runtime path": re.compile(r"/(?:mnt|tmp|Users)/|[A-Za-z]:\\", re.I),
}


def normalize_heading(text: str) -> str:
    text = re.sub(r"^\s*(?:section\s+)?\d+(?:\.\d+)*[.):-]?\s*", "", text, flags=re.I)
    return re.sub(r"\s+", " ", text).strip().lower()


def heading_level(paragraph: object) -> int | None:
    style = getattr(paragraph, "style", None)
    name = getattr(style, "name", "") or ""
    match = re.fullmatch(r"Heading\s+([1-9])", name, re.I)
    return int(match.group(1)) if match else None


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def has_toc_and_update(path: Path) -> tuple[bool, bool]:
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        settings_xml = archive.read("word/settings.xml").decode("utf-8", errors="replace")
    has_toc = bool(re.search(r"<w:instrText[^>]*>[^<]*\bTOC\b", document_xml, re.I))
    update = bool(re.search(r"<w:updateFields\b[^>]*", settings_xml, re.I))
    return has_toc, update


def table_index(document: Document, required_headers: set[str]) -> tuple[int, dict[str, int]] | None:
    for table_number, table in enumerate(document.tables):
        if not table.rows:
            continue
        headers = {
            re.sub(r"\s+", " ", cell.text).strip().lower(): index
            for index, cell in enumerate(table.rows[0].cells)
        }
        if required_headers.issubset(headers):
            return table_number, headers
    return None


def validate_milestones(document: Document, failures: list[str]) -> int:
    required = {
        "milestone id",
        "phase",
        "description",
        "deliverables due",
        "completion criteria",
        "timing or sequence",
        "payment type",
    }
    found = table_index(document, required)
    if found is None:
        failures.append("missing Milestone Schedule table with required headers")
        return 0
    table_number, headers = found
    table = document.tables[table_number]
    if len(table.rows) < 2:
        failures.append("Milestone Schedule has no milestone rows")
        return 0
    required_values = [
        "milestone id",
        "phase",
        "description",
        "deliverables due",
        "completion criteria",
        "payment type",
    ]
    for row_number, row in enumerate(table.rows[1:], start=2):
        for header in required_values:
            value = row.cells[headers[header]].text.strip()
            if not value:
                failures.append(f"Milestone Schedule row {row_number} has blank {header}")
    return len(table.rows) - 1


def validate_deliverables(document: Document, failures: list[str]) -> int:
    required = {"id", "title", "format", "due trigger", "acceptance criteria", "rights category"}
    found = table_index(document, required)
    if found is None:
        failures.append("missing Deliverables table with required headers")
        return 0
    table_number, headers = found
    table = document.tables[table_number]
    if len(table.rows) < 2:
        failures.append("Deliverables table has no deliverable rows")
        return 0
    for row_number, row in enumerate(table.rows[1:], start=2):
        for header in ("id", "title", "due trigger", "acceptance criteria", "rights category"):
            if not row.cells[headers[header]].text.strip():
                failures.append(f"Deliverables row {row_number} has blank {header}")
    return len(table.rows) - 1


def validate_tbd_closeout(document: Document, text: str, failures: list[str]) -> None:
    if "[TBD]" not in text:
        return
    required = {"assumption or constraint", "basis", "owner", "closeout trigger"}
    found = table_index(document, required)
    if found is None:
        failures.append("document contains [TBD] but has no Constraints closeout table")
        return
    table_number, headers = found
    table = document.tables[table_number]
    closeout_rows = 0
    for row_number, row in enumerate(table.rows[1:], start=2):
        assumption = row.cells[headers["assumption or constraint"]].text
        if "[TBD]" not in assumption:
            continue
        closeout_rows += 1
        if not row.cells[headers["owner"]].text.strip():
            failures.append(f"Constraints row {row_number} has [TBD] but no owner")
        if not row.cells[headers["closeout trigger"]].text.strip():
            failures.append(f"Constraints row {row_number} has [TBD] but no closeout trigger")
    if closeout_rows == 0:
        failures.append("document contains [TBD] without a matching Constraints closeout row")


def validate(path: Path) -> dict[str, object]:
    failures: list[str] = []
    try:
        if not zipfile.is_zipfile(path):
            return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            for required in ("[Content_Types].xml", "word/document.xml", "word/styles.xml"):
                if required not in names:
                    failures.append(f"missing DOCX part: {required}")
        document = Document(path)
    except (OSError, ValueError, zipfile.BadZipFile) as exc:
        return {"status": "fail", "failures": [f"cannot read DOCX: {exc}"]}

    text = document_text(document)
    for label, pattern in FORBIDDEN.items():
        if pattern.search(text):
            failures.append(f"document contains forbidden {label}")

    h1 = [
        normalize_heading(paragraph.text)
        for paragraph in document.paragraphs
        if heading_level(paragraph) == 1 and paragraph.text.strip()
    ]
    positions: list[int] = []
    for expected in CORE_SECTIONS:
        candidates = [index for index, value in enumerate(h1) if value == expected]
        if not candidates:
            failures.append(f"missing Heading 1 section: {expected}")
        else:
            positions.append(candidates[0])
    if positions and positions != sorted(positions):
        failures.append("core Heading 1 sections are out of order")

    contribution_positions = [i for i, value in enumerate(h1) if value == "contribution arrangement"]
    follow_on_positions = [i for i, value in enumerate(h1) if value == "production follow-on provisions"]
    constraints_positions = [i for i, value in enumerate(h1) if value == "constraints and assumptions"]
    if constraints_positions:
        end = constraints_positions[0]
        if contribution_positions and contribution_positions[0] > end:
            failures.append("Contribution Arrangement appears after Constraints and Assumptions")
        if follow_on_positions and follow_on_positions[0] > end:
            failures.append("Production Follow-On Provisions appears after Constraints and Assumptions")
        if contribution_positions and follow_on_positions and contribution_positions[0] > follow_on_positions[0]:
            failures.append("Production Follow-On Provisions appears before Contribution Arrangement")

    milestone_count = validate_milestones(document, failures)
    deliverable_count = validate_deliverables(document, failures)
    validate_tbd_closeout(document, text, failures)

    if len(h1) > 8:
        try:
            has_toc, update = has_toc_and_update(path)
        except (KeyError, OSError, zipfile.BadZipFile) as exc:
            failures.append(f"cannot audit TOC settings: {exc}")
        else:
            if not has_toc:
                failures.append("document with more than eight sections has no dynamic TOC field")
            if not update:
                failures.append("document does not set updateFields-on-open")

    return {
        "status": "pass" if not failures else "fail",
        "heading_1_count": len(h1),
        "table_count": len(document.tables),
        "milestone_count": milestone_count,
        "deliverable_count": deliverable_count,
        "failures": failures,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate an OT project-description DOCX.")
    parser.add_argument("document", type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    if not args.document.is_file():
        print(f"ERROR: document not found: {args.document}", file=sys.stderr)
        return 2
    result = validate(args.document)
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("OT project-description structure and separation passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
