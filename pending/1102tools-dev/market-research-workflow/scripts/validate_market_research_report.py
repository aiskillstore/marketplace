#!/usr/bin/env python3
"""Validate a generated Market Research DOCX and its numeric evidence."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


REQUIRED_HEADINGS = [
    "Executive Summary",
    "Requirement and Decision Context",
    "Documents Reviewed",
    "Research Scope and Method",
    "Federal Market Evidence",
    "Commercial and Other Market Evidence",
    "Small-Business and Competition Evidence",
    "Pricing and Contract-Structure Context",
    "Findings and Approved Decisions",
    "Limitations, Conflicts, and Unresolved Questions",
    "Reproducible Search Log",
    "Evidence Register",
]
FORBIDDEN = [
    re.compile(r"\b(?:automatically|therefore)\s+(?:recommend|requires?|proves?)\b", re.I),
    re.compile(r"\b(?:set[- ]aside|commerciality|contract type|price reasonableness)\s+is\s+automatically\b", re.I),
    re.compile(r"\bmcp__|/mnt/|/Users/|[A-Za-z]:\\", re.I),
    re.compile(r"\bghp_[A-Za-z0-9]{20,}\b|\b(?:sk|cfat|SAM)-[A-Za-z0-9_-]{16,}\b", re.I),
]


def all_text(document: Document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def validate(document_path: Path, record_path: Path) -> dict:
    failures: list[str] = []
    if not zipfile.is_zipfile(document_path):
        return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}
    document = Document(document_path)
    record = json.loads(record_path.read_text(encoding="utf-8"))
    text = all_text(document)
    headings = [p.text.strip() for p in document.paragraphs if getattr(p.style, "name", "") == "Heading 1"]
    for heading in REQUIRED_HEADINGS:
        if heading not in headings:
            failures.append(f"missing Heading 1 section: {heading}")
    if [h for h in headings if h in REQUIRED_HEADINGS] != REQUIRED_HEADINGS:
        failures.append("required Heading 1 sections are out of order")
    for pattern in FORBIDDEN:
        if pattern.search(text):
            failures.append(f"forbidden content matched: {pattern.pattern}")
    for item in record.get("findings", []):
        for evidence_id in item.get("evidence_ids", []):
            if evidence_id not in text:
                failures.append(f"finding evidence ID not present in report: {evidence_id}")
    complete = bool(record.get("validation", {}).get("commercial_evidence_complete"))
    if not complete and "Federal-Data Desk-Research Draft" not in text:
        failures.append("incomplete commercial evidence is not labeled as a desk-research draft")
    evidence = [item for item in record.get("evidence", []) if isinstance(item, dict)]
    for index, check in enumerate(record.get("validation", {}).get("numeric_checks", [])):
        expected = sum(float(value) for value in check.get("components", []))
        reported = float(check.get("reported_total", expected))
        label = check.get("label", "numeric check")
        if abs(expected - reported) > 0.005:
            failures.append(f"independent recomputation failed for {label}")
        calculated_total = f"{expected:,.2f}"
        calculation_lines = [
            paragraph.text
            for paragraph in document.paragraphs
            if label in paragraph.text and calculated_total in paragraph.text
        ]
        if not calculation_lines:
            failures.append(f"recomputed total is missing from report: {expected:,.2f}")
            continue
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in evidence
            if item.get("source_class") == "calculation" and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            failures.append(
                f"numeric check {label} does not have exactly one calculation evidence item for {locator}"
            )
            continue
        for evidence_id in calculation_ids:
            if not any(f"[{evidence_id}]" in line for line in calculation_lines):
                failures.append(
                    f"numeric check {label} does not cite its calculation evidence ID: {evidence_id}"
                )
    return {"status": "pass" if not failures else "fail", "heading_count": len(headings), "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path)
    parser.add_argument("--record", required=True, type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    try:
        result = validate(args.document, args.record)
    except (OSError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("Market research DOCX validation passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
