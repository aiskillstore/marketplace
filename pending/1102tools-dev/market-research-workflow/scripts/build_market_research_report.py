#!/usr/bin/env python3
"""Build a structured Market Research DOCX from a validated research record."""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
GREEN = "167D5A"
PALE_GREEN = "EAF4EF"
GRAY = "5B6573"


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_cell_text(cell, value: object, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value if value not in (None, "") else "Not provided"))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12.5, GREEN)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "Evidence ID" not in document.styles:
        style = document.styles.add_style("Evidence ID", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Aptos"
        style.font.size = Pt(8.5)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN)

    header = section.header.paragraphs[0]
    header.text = "1102tools  |  Market Research"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Prepared as of the date shown  |  Page ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_table(document: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[index], NAVY)
        if widths:
            table.rows[0].cells[index].width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if row_index % 2:
                shade(cells[index], "F4F6F8")
            if widths:
                cells[index].width = Inches(widths[index])
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def add_bullets(document: Document, items: list[object], empty: str = "None recorded") -> None:
    if not items:
        document.add_paragraph(empty)
        return
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("decision") or item.get("question") or json.dumps(item, sort_keys=True)
        else:
            text = str(item)
        paragraph = document.add_paragraph(text, style="List Bullet")
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.05


def cite_ids(paragraph, ids: list[str]) -> None:
    if not ids:
        return
    paragraph.add_run(" [")
    run = paragraph.add_run(", ".join(ids), style="Evidence ID")
    run.bold = True
    paragraph.add_run("]")


def build(record: dict, output: Path) -> None:
    validation = record.get("validation", {})
    if record.get("schema_version") != "1.2":
        raise ValueError("market research records must be migrated to schema 1.2 before report generation")
    for field in ("findings_approved", "decisions_approved", "unresolved_items_disposition_approved"):
        if validation.get(field) is not True:
            raise ValueError(f"{field} must be true before report generation")
    document = Document()
    configure(document)
    complete = bool(validation.get("commercial_evidence_complete"))
    title = "FAR Part 10 Market Research Report" if complete else "Federal-Data Desk-Research Draft"
    subtitle = record.get("question", "Market research")
    as_of = record.get("scope", {}).get("as_of_date", "Not stated")

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(115)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("1102tools")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(GREEN)
    title_p = document.add_paragraph(title, style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = document.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    meta = document.add_paragraph(f"As of {as_of}\nPrepared by: ____________________")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    sections = [
        "Executive Summary",
        "Requirement and Decision Context",
        "Documents Reviewed",
        "Research Scope and Method",
        "Federal Market Evidence",
        "Commercial and Other Market Evidence",
        "Small-Business and Competition Evidence",
        "Pricing and Contract-Structure Context",
        "Findings and Approved Decisions",
        "Limitations, Conflicts, and Unresolved Questions",
        "Reproducible Search Log",
        "Evidence Register",
    ]

    document.add_heading(sections[0], level=1)
    document.add_paragraph(record.get("validation", {}).get("executive_summary", "This report organizes the approved research record and its limitations."))
    if not complete:
        note = document.add_paragraph()
        note.style = document.styles["Intense Quote"]
        note.add_run("Completion boundary: ").bold = True
        note.add_run("Commercial-market evidence was not marked complete. This document is a federal-data desk-research draft, not a complete contract-file-ready market research report.")

    document.add_heading(sections[1], level=1)
    document.add_paragraph(record.get("question", "Not provided"))
    scope = record.get("scope", {})
    add_table(document, ["Scope field", "Value"], [[key.replace("_", " ").title(), value] for key, value in scope.items()], [2.1, 4.8])
    document.add_heading("User context and assumptions", level=2)
    add_bullets(document, record.get("user_context", []) + record.get("assumptions", []))

    document.add_heading(sections[2], level=1)
    docs = record.get("document_register", [])
    add_table(
        document,
        ["File", "Type and status", "Role", "Controlling location", "Gaps or conflicts"],
        [[d.get("file", ""), f"{d.get('document_type', '')} / {d.get('status', 'unclear')}", d.get("role", ""), d.get("controlling_location", ""), d.get("gaps_or_conflicts", "")] for d in docs],
        [1.25, 1.25, 1.65, 1.3, 1.45],
    ) if docs else document.add_paragraph("No acquisition documents were available for this research record.")

    document.add_heading(sections[3], level=1)
    document.add_paragraph(record.get("validation", {}).get("methodology", "Sources, scope, and limitations are recorded in the query and evidence registers."))
    document.add_paragraph("Government-wide and agency-specific results are analyzed separately unless an approved method states otherwise.")

    evidence = {item["id"]: item for item in record.get("evidence", []) if isinstance(item, dict) and "id" in item}
    findings = record.get("findings", [])
    buckets = {
        sections[4]: {"federal_mcp"},
        sections[5]: {"official_web", "other_web"},
    }
    for heading, classes in buckets.items():
        document.add_heading(heading, level=1)
        matched = [f for f in findings if any(evidence.get(eid, {}).get("source_class") in classes for eid in f.get("evidence_ids", []))]
        if not matched:
            document.add_paragraph("No findings in this source class were recorded.")
        for finding in matched:
            p = document.add_paragraph(finding.get("text", ""))
            cite_ids(p, finding.get("evidence_ids", []))

    document.add_heading(sections[6], level=1)
    document.add_paragraph(record.get("validation", {}).get("small_business_analysis", "No approved small-business or competition analysis was recorded."))
    document.add_paragraph("Historical award percentages inform research but do not by themselves establish the FAR 19.502-2 Rule of Two or a set-aside decision.")

    document.add_heading(sections[7], level=1)
    document.add_paragraph(record.get("validation", {}).get("pricing_analysis", "No approved pricing or contract-structure analysis was recorded."))
    for index, check in enumerate(record.get("validation", {}).get("numeric_checks", [])):
        components = [float(value) for value in check.get("components", [])]
        locator = f"validation.numeric_checks[{index}]"
        calculation_ids = [
            item.get("id")
            for item in record.get("evidence", [])
            if isinstance(item, dict)
            and item.get("source_class") == "calculation"
            and item.get("locator") == locator
        ]
        if len(calculation_ids) != 1:
            raise ValueError(
                f"numeric check {index} requires exactly one calculation evidence item whose locator is {locator}"
            )
        paragraph = document.add_paragraph(f"{check.get('label', 'Calculated total')}: {sum(components):,.2f}")
        cite_ids(paragraph, calculation_ids)

    document.add_heading(sections[8], level=1)
    for finding in findings:
        p = document.add_paragraph(finding.get("text", ""), style="List Bullet")
        cite_ids(p, finding.get("evidence_ids", []))
    document.add_heading("Approved user decisions", level=2)
    add_bullets(document, record.get("user_decisions", []), "No acquisition decision is recorded as approved.")

    document.add_heading(sections[9], level=1)
    add_bullets(document, record.get("conflicts", []), "No unresolved source conflict was recorded.")
    add_bullets(document, record.get("unresolved_questions", []), "No unresolved question was recorded.")
    for item in record.get("inferences", []):
        p = document.add_paragraph("Inference: " + item.get("text", item.get("reasoning", "")), style="List Bullet")
        cite_ids(p, item.get("evidence_ids", []))

    document.add_heading(sections[10], level=1)
    queries = record.get("queries", [])
    add_table(
        document,
        ["Source / operation", "Sanitized parameters", "Retrieved", "Coverage and limits"],
        [[q.get("operation", q.get("source", "")), json.dumps(q.get("parameters", {}), sort_keys=True), q.get("retrieved_at", ""), f"{q.get('count', 'n/a')}; {q.get('limitations', '')}"] for q in queries],
        [1.5, 2.35, 1.25, 2.0],
    ) if queries else document.add_paragraph("No external query was made.")

    document.add_heading(sections[11], level=1)
    add_table(
        document,
        ["ID", "Class", "Source", "Fact", "Limitations"],
        [[e.get("id", ""), e.get("source_class", ""), f"{e.get('title', '')}\n{e.get('locator', '')}", e.get("fact", ""), e.get("limitations", "")] for e in record.get("evidence", [])],
        [0.55, 0.85, 1.55, 2.5, 1.65],
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("record", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    validator = Path(__file__).with_name("validate_research_record.py")
    result = subprocess.run([sys.executable, str(validator), str(args.record)], check=False)
    if result.returncode:
        return result.returncode
    record = json.loads(args.record.read_text(encoding="utf-8"))
    build(record, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
