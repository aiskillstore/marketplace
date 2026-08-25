#!/usr/bin/env python3
"""Validate a generated Acquisition Policy Impact Brief."""

from __future__ import annotations

import argparse
import importlib.util
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


REQUIRED_HEADINGS = [
    "Executive Summary",
    "Question and Scope",
    "Documented Current Status",
    "Source Hierarchy and Authorities",
    "Change Timeline",
    "Government and Industry Impacts",
    "Open Issues and Comment Deadlines",
    "Operational Considerations",
    "Evidence Register",
    "Limitations and Reserved Determinations",
]
FORBIDDEN = [
    re.compile(r"\bmcp__|/mnt/|/Users/|[A-Za-z]:\\", re.I),
    re.compile(r"\bgh[pousr]_[A-Za-z0-9]{20,}\b|\b(?:sk|cfat|SAM)-[A-Za-z0-9_-]{16,}\b", re.I),
    re.compile(r"\b(?:this model deviation|model deviation text)\s+(?:is|was)\s+(?:legally )?(?:operative|applicable)\b", re.I),
]
BOUNDARY_PHRASES = (
    "authorized agency official must determine procurement-specific applicability",
    "does not provide legal advice",
)


def load_record_validator(path: Path):
    spec = importlib.util.spec_from_file_location("policy_record_validator", path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def external_hyperlinks(document: Document) -> set[str]:
    urls: set[str] = set()
    for relationship in document.part.rels.values():
        if relationship.reltype.endswith("/hyperlink") and relationship.is_external:
            urls.add(relationship.target_ref)
    return urls


def table_geometry_failures(document: Document) -> list[str]:
    failures: list[str] = []
    for table_index, table in enumerate(document.tables):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa" or tbl_w.get(qn("w:w")) != "9360":
            failures.append(f"table {table_index + 1} does not have fixed 9360-DXA width")
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != "120":
            failures.append(f"table {table_index + 1} does not have 120-DXA indent")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            failures.append(f"table {table_index + 1} does not use fixed layout")
        grid_widths = [int(col.get(qn("w:w"), "0")) for col in table._tbl.tblGrid]
        if not grid_widths or sum(grid_widths) != 9360:
            failures.append(f"table {table_index + 1} grid widths do not total 9360 DXA")
        header_props = table.rows[0]._tr.get_or_add_trPr()
        if header_props.find(qn("w:tblHeader")) is None:
            failures.append(f"table {table_index + 1} does not repeat its header row")
        for row_index, row in enumerate(table.rows):
            widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                widths.append(int(tc_w.get(qn("w:w"), "0")) if tc_w is not None else 0)
            if grid_widths and widths != grid_widths:
                failures.append(f"table {table_index + 1} row {row_index + 1} cell widths differ from grid")
                break
    return failures


def validate(document_path: Path, record_path: Path) -> dict:
    failures: list[str] = []
    if not zipfile.is_zipfile(document_path):
        return {"status": "fail", "failures": ["file is not a valid DOCX ZIP"]}

    record = json.loads(record_path.read_text(encoding="utf-8"))
    record_validator = load_record_validator(Path(__file__).with_name("validate_policy_research_record.py"))
    record_result = record_validator.validate_record(record)
    if record_result["status"] != "pass":
        failures.extend(f"record: {failure}" for failure in record_result["failures"])

    document = Document(document_path)
    text = all_text(document)
    headings = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if getattr(paragraph.style, "name", "") == "Heading 1"
    ]
    for heading in REQUIRED_HEADINGS:
        if heading not in headings:
            failures.append(f"missing Heading 1 section: {heading}")
    if [heading for heading in headings if heading in REQUIRED_HEADINGS] != REQUIRED_HEADINGS:
        failures.append("required Heading 1 sections are out of order")

    for pattern in FORBIDDEN:
        if pattern.search(text):
            failures.append(f"forbidden content matched: {pattern.pattern}")
    lowered = text.lower()
    for phrase in BOUNDARY_PHRASES:
        if phrase not in lowered:
            failures.append(f"required decision-boundary language is missing: {phrase}")
    as_of = record.get("scope", {}).get("as_of_date", "")
    if as_of and as_of not in text:
        failures.append("record as-of date is missing from the brief")

    evidence_ids = {item.get("id") for item in record.get("evidence", []) if isinstance(item, dict)}
    for finding in record.get("findings", []):
        for evidence_id in finding.get("evidence_ids", []):
            if evidence_id not in text:
                failures.append(f"finding evidence ID not present in brief: {evidence_id}")
    for policy in record.get("policy_items", []):
        policy_id = policy.get("id")
        if policy_id and policy_id not in text:
            failures.append(f"policy item ID not present in brief: {policy_id}")
    for evidence_id in sorted(evidence_ids):
        if evidence_id not in text:
            failures.append(f"evidence register is missing ID: {evidence_id}")

    urls = external_hyperlinks(document)
    for item in record.get("evidence", []):
        url = item.get("canonical_url", "")
        if url and url not in urls:
            failures.append(f"evidence URL is not a live DOCX hyperlink: {url}")

    failures.extend(table_geometry_failures(document))
    if len(document.tables) < 3:
        failures.append("brief must contain at least three structured evidence tables")

    return {
        "status": "pass" if not failures else "fail",
        "heading_count": len(headings),
        "table_count": len(document.tables),
        "hyperlink_count": len(urls),
        "failures": failures,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path)
    parser.add_argument("--record", required=True, type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    try:
        result = validate(args.document, args.record)
    except (OSError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    if args.json:
        print(json.dumps(result, indent=2, sort_keys=True))
    elif result["status"] == "pass":
        print("Acquisition Policy DOCX validation passed.")
    else:
        print("VALIDATION FAILED")
        for failure in result["failures"]:
            print(f"- {failure}")
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
