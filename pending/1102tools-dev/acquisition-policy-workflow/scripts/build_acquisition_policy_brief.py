#!/usr/bin/env python3
"""Build a validated Acquisition Policy Impact Brief from a policy record."""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "5B6573"
BLACK = "202124"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, *, name: str = "Calibri", size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str) -> None:
    if not url:
        return
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Policy Title" not in document.styles:
        title = document.styles.add_style("Policy Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = document.styles["Policy Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(BLACK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True

    if "Evidence ID" not in document.styles:
        evidence_style = document.styles.add_style("Evidence ID", WD_STYLE_TYPE.CHARACTER)
    else:
        evidence_style = document.styles["Evidence ID"]
    evidence_style.font.name = "Calibri"
    evidence_style.font.size = Pt(9)
    evidence_style.font.bold = True
    evidence_style.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    header = section.header.paragraphs[0]
    header.text = "1102tools  |  Acquisition Policy Impact Brief"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=8, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer.add_run("Prepared as of the date shown  |  Page ")
    set_run_font(prefix, size=8, color=MID_GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_metadata(document: Document, label: str, value: object) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True)
    value_run = paragraph.add_run(str(value if value not in (None, "") else "Not provided"))
    set_run_font(value_run)


def set_cell_text(cell, value: object, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(value if value not in (None, "") else "Not stated"))
    set_run_font(run, size=9.5, color=color, bold=bold)


def add_table(document: Document, headers: list[str], rows: list[list[object]], widths_dxa: list[int]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=WHITE)
        shade(table.rows[0].cells[index], DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
            if row_index % 2:
                shade(cells[index], LIGHT_GRAY)
    set_table_geometry(table, widths_dxa)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_bullets(document: Document, items: list[object], empty_text: str) -> None:
    if not items:
        document.add_paragraph(empty_text)
        return
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("question") or item.get("description") or json.dumps(item, sort_keys=True)
            evidence_ids = item.get("evidence_ids", [])
        else:
            text = str(item)
            evidence_ids = []
        paragraph = document.add_paragraph(str(text), style="List Bullet")
        cite_ids(paragraph, evidence_ids)


def cite_ids(paragraph, ids: list[str]) -> None:
    if not ids:
        return
    paragraph.add_run(" [")
    paragraph.add_run(", ".join(ids), style="Evidence ID")
    paragraph.add_run("]")


def display_value(value: object) -> str:
    if value in (None, "", [], {}):
        return "Not stated"
    if isinstance(value, dict):
        value = {key: item for key, item in value.items() if item not in (None, "", [], {})}
        if not value:
            return "Not stated"
        return json.dumps(value, sort_keys=True)
    if isinstance(value, list):
        return json.dumps(value, sort_keys=True)
    return str(value)


def policy_rows(record: dict) -> list[list[object]]:
    return [
        [
            item.get("id", ""),
            item.get("status", "").replace("_", " ").title(),
            item.get("citation", ""),
            item.get("agency", "") or "Government-wide or not stated",
            item.get("applicability_summary", ""),
            ", ".join(item.get("evidence_ids", [])),
        ]
        for item in record.get("policy_items", [])
    ]


def impact_items(record: dict, lens: str) -> list[dict]:
    impacts = record.get("validation", {}).get("impacts", {})
    if not isinstance(impacts, dict):
        return []
    if lens == "neutral":
        return list(impacts.get("government", [])) + list(impacts.get("industry", []))
    return list(impacts.get(lens, []))


def build(record: dict, output: Path) -> None:
    validation = record.get("validation", {})
    if not validation.get("findings_approved") or not validation.get("brief_approved"):
        raise ValueError("findings and brief generation must be approved before building the DOCX")

    document = Document()
    configure_styles(document)
    request = record["request"]
    scope = record["scope"]
    as_of = scope["as_of_date"]
    agency = scope.get("agency") or "Published federal acquisition policy"
    lens = request["audience_lens"]

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("ACQUISITION POLICY")
    set_run_font(run, size=10, color=BLUE, bold=True)
    title = document.add_paragraph("Acquisition Policy Impact Brief", style="Policy Title")
    subtitle = document.add_paragraph(request.get("question", ""))
    subtitle.paragraph_format.space_after = Pt(14)
    if subtitle.runs:
        set_run_font(subtitle.runs[0], size=13, color=MID_GRAY)
    add_metadata(document, "Agency or scope", agency)
    add_metadata(document, "As of", as_of)
    add_metadata(document, "Audience lens", lens.title())
    add_metadata(document, "Status", "Documented published-source analysis; not a legal opinion or procurement-specific determination")
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(validation.get("executive_summary", "No approved executive summary was supplied."))
    boundary = document.add_paragraph(style="Intense Quote")
    boundary.add_run("Decision boundary: ").bold = True
    boundary.add_run(
        "This brief states what cited published sources indicate as of the date shown. "
        "An authorized agency official must determine procurement-specific applicability."
    )

    document.add_heading("Question and Scope", level=1)
    document.add_paragraph(request.get("question", "Not stated"))
    scope_rows = [[key.replace("_", " ").title(), display_value(value)] for key, value in scope.items()]
    add_table(document, ["Scope field", "Approved value"], scope_rows, [2700, 6660])

    document.add_heading("Documented Current Status", level=1)
    if record.get("policy_items"):
        add_table(
            document,
            ["ID", "Status", "Citation", "Agency", "Documented treatment", "Evidence"],
            policy_rows(record),
            [850, 1450, 1200, 1450, 3060, 1350],
        )
    else:
        document.add_paragraph("No approved policy item was recorded.")
    document.add_heading("Approved findings", level=2)
    for finding in record.get("findings", []):
        paragraph = document.add_paragraph(finding.get("text", ""), style="List Bullet")
        cite_ids(paragraph, finding.get("evidence_ids", []))
    if not record.get("findings"):
        document.add_paragraph("No approved finding was recorded.")

    document.add_heading("Source Hierarchy and Authorities", level=1)
    hierarchy = validation.get("source_hierarchy", [])
    add_bullets(
        document,
        hierarchy,
        "The analysis distinguishes codified text, agency deviations, model text, rulemaking, guidance, comments, and supplied documents.",
    )

    document.add_heading("Change Timeline", level=1)
    timeline = record.get("timeline", [])
    if timeline:
        add_table(
            document,
            ["Date", "Event", "Status", "Evidence"],
            [[item.get("date", ""), item.get("event", ""), item.get("status", ""), ", ".join(item.get("evidence_ids", []))] for item in timeline],
            [1250, 4700, 1850, 1560],
        )
    else:
        document.add_paragraph("No change event was required for the approved scope.")

    document.add_heading("Government and Industry Impacts", level=1)
    if lens == "neutral":
        document.add_heading("Government lens", level=2)
        add_bullets(document, validation.get("impacts", {}).get("government", []), "No approved government impact was recorded.")
        document.add_heading("Industry lens", level=2)
        add_bullets(document, validation.get("impacts", {}).get("industry", []), "No approved industry impact was recorded.")
    else:
        document.add_heading(f"{lens.title()} lens", level=2)
        add_bullets(document, impact_items(record, lens), f"No approved {lens} impact was recorded.")

    document.add_heading("Open Issues and Comment Deadlines", level=1)
    deadlines = validation.get("open_issues", [])
    add_bullets(document, deadlines, "No open issue or comment deadline was identified within the approved scope.")
    if record.get("stakeholder_positions"):
        document.add_heading("Observed stakeholder positions", level=2)
        add_table(
            document,
            ["Submitter type", "Observed position", "Sample", "Coverage and limitations", "Evidence"],
            [
                [
                    item.get("submitter_type", ""),
                    item.get("position", ""),
                    f"{item.get('reviewed_count', 0)} of {item.get('returned_count', 0)}; {item.get('sample_method', '')}",
                    item.get("limitations", ""),
                    ", ".join(item.get("evidence_ids", [])),
                ]
                for item in record["stakeholder_positions"]
            ],
            [1250, 3000, 1600, 2350, 1160],
        )

    document.add_heading("Operational Considerations", level=1)
    add_bullets(document, validation.get("operational_considerations", []), "No operational consideration was approved.")
    if record.get("conflicts"):
        document.add_heading("Conflicts", level=2)
        add_table(
            document,
            ["ID", "Issue", "Status", "Resolution and source"],
            [
                [
                    item.get("id", ""),
                    item.get("issue", ""),
                    item.get("status", "").replace("_", " ").title(),
                    (
                        f"{item.get('resolution', '')} ({item.get('resolved_by', '')}; {item.get('resolved_at', '')})"
                        if item.get("resolution")
                        else "Reserved to an authorized official"
                    ),
                ]
                for item in record["conflicts"]
            ],
            [850, 4200, 1700, 2610],
        )
    if record.get("unresolved_questions"):
        document.add_heading("Unresolved questions", level=2)
        add_bullets(document, record["unresolved_questions"], "")

    document.add_page_break()
    document.add_heading("Evidence Register", level=1)
    evidence = record.get("evidence", [])
    if evidence:
        table = add_table(
            document,
            ["ID", "Type", "Source and locator", "Supported fact", "Limits"],
            [
                [
                    item.get("id", ""),
                    item.get("source_type", "").replace("_", " ").title(),
                    f"{item.get('title', '')}\n{item.get('locator', '')}",
                    item.get("fact", ""),
                    item.get("limitations", ""),
                ]
                for item in evidence
            ],
            [850, 1400, 2350, 3000, 1760],
        )
        for row, item in zip(table.rows[1:], evidence):
            url = item.get("canonical_url", "")
            if url:
                paragraph = row.cells[2].add_paragraph()
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(0)
                add_hyperlink(paragraph, urlparse(url).netloc or "Official source", url)
    else:
        document.add_paragraph("No evidence item was recorded.")

    document.add_heading("Limitations and Reserved Determinations", level=1)
    add_bullets(document, record.get("limitations", []), "No additional limitation was recorded.")
    document.add_paragraph(
        "This brief does not provide legal advice, approve policy, select clauses, or determine which rule legally governs a specific procurement. "
        "Confirm transaction-specific treatment with the responsible contracting, policy, and legal officials."
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("record", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    validator = Path(__file__).with_name("validate_policy_research_record.py")
    result = subprocess.run([sys.executable, str(validator), str(args.record)], check=False)
    if result.returncode:
        return result.returncode
    try:
        build(json.loads(args.record.read_text(encoding="utf-8")), args.output)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
